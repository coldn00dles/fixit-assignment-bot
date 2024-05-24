import docx
from openai import OpenAI
import pandas as pd
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from dotenv import load_dotenv
import os


load_dotenv()
client = OpenAI(
    api_key= os.getenv("OPENAI_API_KEY")
)

def textprocessing(file_path):
    """
    Takes a document and parses it to seperate normal text from tables. \n
    Returns a tuple of text(List) and table(List)"""
    doc = docx.Document(file_path)
    text = []
    tables = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  
            text.append(paragraph.text)

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)
        tables.append(table_data)

    return text, tables

def embed_query(docs):
    """
    Takes a document and returns vector embeddings for it.
    """
    embds = []
    for text in docs:
        response = client.embeddings.create(input=text, model="text-embedding-ada-002")
        embds.append(response.data[0].embedding)
    return embds

def setvecdb(text_embds, table_embds, text, tables):
    """
    Takes embeddings and normal documents for both tables and texts and makes a simplified collection out of it \n
    Returns a Pandas Dataframe with the document content, type of document, and embedding as columns"""
    data = []
    
    for idx, embedding in enumerate(text_embds):
        data.append({"type": "text", "content": text[idx], "embedding": embedding})

    for idx, (embedding, table_text) in enumerate(zip(table_embds, tables)):
        data.append({"type": "table", "content": table_text, "embedding": embedding})
    
    vectordb = pd.DataFrame(data)
    return vectordb

def retriever(vectordb, query_text):
    """
    Takes a Pandas Dataframe as the vector database and a question(string) \n
    Compares cosine similarity between the vectorised question and embeddings in the dataframe, returns the top 4 text/tabular data contents ordered by the similiarity as a DataFrame"""
    qembed = client.embeddings.create(input=query_text, model="text-embedding-ada-002")
    query_embedding = np.array(qembed.data[0].embedding).reshape(1, -1)
    
    embeddings = np.stack(vectordb['embedding'].values)
    similarities = cosine_similarity(query_embedding, embeddings).flatten()
    
    vectordb['similarity'] = similarities
    results = vectordb.nlargest(4, 'similarity')
    
    return results

prompt = """You are a smart chatbot that answers simple and complex questions of a user on the basis of contents from a document uploaded by the user. The document can have tabular data and textual data, with tables seperated by bar(|) and hyphens(-) to represent its rows and columns.
Don't simply rely on direct logic for questions and use complex reasoning often for answering them. Below given is the context and the question. If unaware of the answer, simply respond that you cannot answer it based on the data, do not hallucinate. Make sure to answer questions related to tabular data properly as well. 
Explain your purpose if asked, make sure you are helpful"""

def generate_answer(question,ctxt):
    """
    Takes two strings : the user's question and the context retrieved \n
    LLM generates appropiate answer and returns a string"""
    response =  client.chat.completions.create(
        model = "gpt-3.5-turbo-0125",
        messages = [
            {
                "role" : "system",
                "content" : prompt
            },
            {"role" : "user","content" : question},
            {"role" : "assistant","content" : ctxt}
        ]
    )
    return response.choices[0].message.content